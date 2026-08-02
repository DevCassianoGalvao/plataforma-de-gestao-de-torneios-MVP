from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from xml.sax.saxutils import escape


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'docs' / 'MANUAL_DO_USUARIO.md'
OUTPUT = ROOT / 'docs' / 'Manual_do_Usuario_Torneio_Online_Web_App.docx'
PDF_OUTPUT = ROOT / 'docs' / 'Manual_do_Usuario_Torneio_Online_Web_App.pdf'

NAVY = RGBColor(9, 23, 38)
GREEN = RGBColor(18, 156, 97)
MUTED = RGBColor(82, 98, 112)


def shade(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    properties.append(shading)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def add_text(paragraph, text: str, bold=False, color=None):
    cursor = 0
    for match in re.finditer(r'\*\*(.+?)\*\*', text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        if color:
            run.font.color.rgb = color
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.bold = bold
        if color:
            run.font.color.rgb = color


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.14

    for style_name, size, color, before, after in [
        ('Title', 27, NAVY, 0, 8),
        ('Heading 1', 18, NAVY, 18, 8),
        ('Heading 2', 13.5, GREEN, 13, 6),
        ('Heading 3', 11.5, NAVY, 10, 4),
    ]:
        style = styles[style_name]
        style.font.name = 'Arial'
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Torneio Online Web App | Manual do Usuário')
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_cover(document: Document, title: str, subtitle: str, version: str) -> None:
    section = document.sections[0]
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(110)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(29)
    run.font.bold = True
    run.font.color.rgb = NAVY
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Torneio Online Web App')
    run.font.name = 'Arial'
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = GREEN
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(45)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(version)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    document.add_page_break()


def add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    shade(paragraph, 'E5F5ED')
    add_text(paragraph, text, bold=False, color=NAVY)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def pdf_markup(text: str) -> str:
    safe = escape(text)
    return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', safe)


def build_pdf(lines: list[str]) -> None:
    document = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        rightMargin=2.0 * cm,
        leftMargin=2.0 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title='Manual do Usuário - Torneio Online Web App',
        author='Torneio Online Web App',
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle('ManualBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=10.2, leading=14, textColor=colors.HexColor('#091726'), spaceAfter=7)
    h1 = ParagraphStyle('ManualH1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=18, leading=22, textColor=colors.HexColor('#091726'), spaceBefore=16, spaceAfter=9)
    h2 = ParagraphStyle('ManualH2', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=13.2, leading=17, textColor=colors.HexColor('#129C61'), spaceBefore=11, spaceAfter=6)
    cover_title = ParagraphStyle('CoverTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=28, leading=33, textColor=colors.HexColor('#091726'), alignment=TA_CENTER)
    cover_brand = ParagraphStyle('CoverBrand', parent=body, fontName='Helvetica-Bold', fontSize=17, leading=22, textColor=colors.HexColor('#129C61'), alignment=TA_CENTER, spaceAfter=12)
    cover_subtitle = ParagraphStyle('CoverSubtitle', parent=body, fontSize=12, leading=17, textColor=colors.HexColor('#526270'), alignment=TA_CENTER)
    story = [Spacer(1, 6.5 * cm), Paragraph('Manual do Usuário', cover_title), Spacer(1, 0.45 * cm), Paragraph('Torneio Online Web App', cover_brand), Paragraph('Guia de operação para campeonatos de futebol', cover_subtitle), Spacer(1, 1.5 * cm), Paragraph('Versão 1.0 - Agosto de 2026', cover_subtitle), PageBreak()]
    ordered_pattern = re.compile(r'^\d+\.\s+(.*)$')
    bullets: list[str] = []

    def flush_bullets() -> None:
        nonlocal bullets
        if bullets:
            story.append(ListFlowable([ListItem(Paragraph(pdf_markup(item), body), leftIndent=9) for item in bullets], bulletType='bullet', leftIndent=16, bulletFontName='Helvetica'))
            story.append(Spacer(1, 3))
            bullets = []

    for line in lines[8:]:
        text = line.strip()
        if not text:
            flush_bullets()
            continue
        if text.startswith('## '):
            flush_bullets()
            story.append(Paragraph(pdf_markup(text[3:]), h1))
        elif text.startswith('### '):
            flush_bullets()
            story.append(Paragraph(pdf_markup(text[4:]), h2))
        elif text.startswith('> '):
            flush_bullets()
            box = Table([[Paragraph(pdf_markup(text[2:]), body)]], colWidths=[16.9 * cm])
            box.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#E5F5ED')), ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#B7E5CF')), ('LEFTPADDING', (0, 0), (-1, -1), 10), ('RIGHTPADDING', (0, 0), (-1, -1), 10), ('TOPPADDING', (0, 0), (-1, -1), 8), ('BOTTOMPADDING', (0, 0), (-1, -1), 8)]))
            story.append(KeepTogether(box))
            story.append(Spacer(1, 5))
        elif text.startswith('- '):
            bullets.append(text[2:])
        else:
            flush_bullets()
            match = ordered_pattern.match(text)
            if match:
                story.append(Paragraph(pdf_markup(match.group(1)), body, bulletText='•'))
            else:
                story.append(Paragraph(pdf_markup(text), body))

    flush_bullets()

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.HexColor('#526270'))
        canvas.drawCentredString(A4[0] / 2, 1.05 * cm, 'Torneio Online Web App | Manual do Usuário')
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)


def build() -> None:
    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    document = Document()
    configure(document)

    title = lines[0].removeprefix('# ').strip()
    # The opening block uses blank lines between the product name, subtitle
    # and version. Keep the cover independent from those visual separators.
    subtitle = lines[4].strip() if len(lines) > 4 else ''
    version = lines[6].strip() if len(lines) > 6 else ''
    add_cover(document, title, subtitle, version)

    ordered_pattern = re.compile(r'^\d+\.\s+(.*)$')
    for line in lines[6:]:
        text = line.strip()
        if not text:
            continue
        if text.startswith('# '):
            continue
        if text.startswith('## '):
            document.add_heading(text[3:], level=1)
            continue
        if text.startswith('### '):
            document.add_heading(text[4:], level=2)
            continue
        if text.startswith('> '):
            add_callout(document, text[2:])
            continue
        if text.startswith('- '):
            paragraph = document.add_paragraph(style='List Bullet')
            add_text(paragraph, text[2:])
            continue
        match = ordered_pattern.match(text)
        if match:
            paragraph = document.add_paragraph(style='List Number')
            add_text(paragraph, match.group(1))
            continue
        paragraph = document.add_paragraph()
        add_text(paragraph, text)

    document.save(OUTPUT)
    build_pdf(lines)
    print(OUTPUT)
    print(PDF_OUTPUT)


if __name__ == '__main__':
    build()
